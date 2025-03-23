import os
import fitz  # PyMuPDF
import spacy
import re
import nltk
import pdfplumber
import textstat
from collections import Counter
from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
import language_tool_python
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from bs4 import BeautifulSoup
import difflib
import google.generativeai as genai
from sentence_transformers import SentenceTransformer, util
import io
from PIL import Image


nlp = spacy.load("en_core_web_trf") 
model = SentenceTransformer("all-MiniLM-L6-v2")

GENAI_API_KEY = "AIzaSyAIeSiWtSZ5-LbGrN69Fveisjxq5ehb2S8"  
genai.configure(api_key=GENAI_API_KEY)


nltk.download('punkt')
nlp = spacy.load("en_core_web_sm")
tool = language_tool_python.LanguageTool('en-US')

app = FastAPI()


HARD_SKILLS = ["Python", "Java", "Machine Learning", "SQL", "TensorFlow", "AWS"]
SOFT_SKILLS = [
    "communication", "leadership", "teamwork", "problem-solving", "adaptability", 
    "creativity", "time management", "empathy", "collaboration", "critical thinking", 
    "negotiation", "organization", "decision making", "emotional intelligence", 
    "conflict resolution", "interpersonal skills", "work ethic", "attention to detail", 
    "flexibility", "customer service", "multitasking", "self-motivation", "patience", 
    "stress management", "active listening", "public speaking", "persuasion", 
    "strategic thinking", "innovation", "networking", "mentoring", "delegation", 
    "presentation skills", "adaptability", "initiative", "dependability", "integrity", 
    "professionalism", "resourcefulness", "goal setting", "analytical skills", 
    "research skills", "writing skills", "editing skills", "planning", "prioritization", 
    "self-discipline", "self-awareness", "cultural awareness", "open-mindedness", 
    "tactfulness", "diplomacy", "negotiation", "mediation", "team building", 
    "relationship management", "influencing", "coaching", "training", "facilitation", 
    "brainstorming", "decision making", "problem sensitivity", "service orientation", 
    "social perceptiveness", "coordination", "persuasion", "instructing", "quality control", 
    "monitoring", "complex problem solving", "judgment and decision making", 
    "systems analysis", "systems evaluation", "time management", "management of financial resources", 
    "management of material resources", "management of personnel resources"
]


ACTION_VERBS = [
    "led", "managed", "developed", "achieved", "designed", "implemented", "created", 
    "initiated", "executed", "coordinated", "supervised", "facilitated", "directed", 
    "produced", "engineered", "constructed", "formulated", "organized", "planned", 
    "administered", "built", "launched", "pioneered", "enhanced", "optimized", 
    "streamlined", "transformed", "upgraded", "improved", "innovated", "devised",
    "analyzed", "assessed", "calculated", "compiled", "computed", "designed", 
    "developed", "evaluated", "examined", "explored", "forecasted", "formulated", 
    "identified", "interpreted", "investigated", "measured", "modeled", "predicted", 
    "researched", "reviewed", "solved", "tested", "validated", "verified", 
    "accelerated", "accomplished", "achieved", "advanced", "boosted", "completed", 
    "delivered", "demonstrated", "drove", "earned", "exceeded", "expanded", 
    "generated", "improved", "increased", "maximized", "outperformed", "produced", 
    "reached", "realized", "reduced", "succeeded", "surpassed"
]

class ATSResponse(BaseModel):
    ats_score: float
    parse_rate: float
    impact: bool
    repeated_words: list
    grammar_issues: list
    format_check: str
    resume_length: str
    long_bullets: list
    contact_info: dict
    sections: dict
    hard_skills: list
    soft_skills: list
    design_type: str
import re

def extract_text(file_path):
    text = ""
    
    # Check file type
    if file_path.endswith(".pdf"):
        try:
            with pdfplumber.open(file_path) as pdf:
                text = " ".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            print("PDF extraction error:", e)

    elif file_path.endswith(".docx"):
        try:
            doc = Document(file_path)
            text = " ".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print("DOCX extraction error:", e)
    
    elif file_path.endswith(".txt"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            print("TXT extraction error:", e)
    
    elif file_path.endswith(".html"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                text = soup.get_text(separator=" ")
        except Exception as e:
            print("HTML extraction error:", e)
    
    # Remove any potential hash-like patterns that might represent images
    text = re.sub(r'\b[a-fA-F0-9]{16,}\b', '', text)  # Removes long hex-like sequences
    
    return text.strip()


# This function gets content from images using OCR
def extract_image_text(file_path):
    ocr_text = ""
    image_count = 0
    
    try:
        # Open the PDF file
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            # Get images from page
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                image_count += 1
                xref = img[0]
                
                # Extract image
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Use PIL to open the image
                image = Image.open(io.BytesIO(image_bytes))
                
                # Skip very small images (likely icons or decorative elements)
                if image.width < 50 or image.height < 50:
                    continue
                    
                # Perform OCR on the image
                try:
                    img_text = pytesseract.image_to_string(image)
                    if img_text.strip():  # Only add non-empty text
                        ocr_text += " " + img_text.strip()
                except Exception as ocr_err:
                    print(f"OCR error on image {img_index} on page {page_num+1}: {ocr_err}")
    
    except Exception as e:
        print(f"Image extraction error: {e}")
    
    return ocr_text.strip(), image_count


# Corrected parse rate calculation that accounts for images
def calculate_parse_rate(file_path, extracted_text):
    # Get text from images using OCR
    ocr_text, image_count = extract_image_text(file_path)
    
    # Clean both texts for better comparison
    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip().lower()
    extracted_text = re.sub(r'[^\w\s]', '', extracted_text)
    
    ocr_text = re.sub(r'\s+', ' ', ocr_text).strip().lower()
    ocr_text = re.sub(r'[^\w\s]', '', ocr_text)
    
    # If no images with text, return high parse rate
    if image_count == 0 or not ocr_text:
        return 95.0  # Default high rate when no images with text detected
    
    # Tokenize texts for comparison
    extracted_words = set(extracted_text.split())
    ocr_words = set(ocr_text.split())
    
    # Find unique words in OCR that weren't in the extracted text
    unique_ocr_words = ocr_words - extracted_words
    
    # Calculate how much content was missed in the extraction
    total_unique_words = len(extracted_words.union(ocr_words))
    
    if total_unique_words == 0:
        return 0.0
    
    # Estimate how much content is potentially missing
    # A higher number of unique words in OCR text indicates more missed content
    missed_percentage = (len(unique_ocr_words) / total_unique_words) * 100
    
    # Calculate parse rate: adjust for the estimated missed content
    parse_rate = 100.0 - min(missed_percentage, 100.0)
    
    return max(0, parse_rate)  # Ensure non-negative


# File format check
def check_file_format(file_path):
    if file_path.endswith(".pdf"):
        return "Valid (PDF format)"
    return "Invalid (Only PDFs are recommended)"

# Quantify impact
def quantify_impact(text):
    doc = nlp(text)
    impact_words = [token.text.lower() for token in doc if token.text.lower() in ACTION_VERBS]
    return bool(impact_words)

from nltk.corpus import stopwords
import string

nltk.download("stopwords")
STOPWORDS = set(stopwords.words("english"))
def check_repetition(text):
    words = nltk.word_tokenize(text.lower())  # Tokenize and lowercase
    words = [word for word in words if word not in STOPWORDS and word not in string.punctuation and word.isalpha()]  # Remove stopwords, punctuation, and non-alphabetic tokens
    word_freq = Counter(words)
    repeated = [word for word, count in word_freq.items() if count > 3]  # Words appearing >3 times
    return repeated if repeated else "No excessive repetition"


# Grammar & spelling check (return incorrect & corrected lines)
TECHNICAL_TERMS = {"TensorFlow", "PyTorch", "ChexNet", "AlexNet", "Colab"}  # Add more as needed

nlp = spacy.load("en_core_web_sm")
tool = language_tool_python.LanguageTool('en-US')

# Define technical terms (ensure it's properly initialized)
TECHNICAL_TERMS = {"AI", "ML", "TensorFlow", "PyTorch", "API", "CID"}  # Add more if needed

def check_grammar(text):
    doc = nlp(text)  # Run NLP pipeline
    
    # Detect proper nouns (names, organizations, events, products)
    proper_nouns = {ent.text.lower() for ent in doc.ents if ent.label_ in ["PERSON", "ORG", "EVENT", "PRODUCT"]}
    excluded_terms = proper_nouns.union({term.lower() for term in TECHNICAL_TERMS})  # Combine sets
    
    matches = tool.check(text)
    grammar_issues = []

    for match in matches:
        incorrect = text[match.offset: match.offset + match.errorLength]
        
        # Skip correction if it's a proper noun, technical term, or special character
        if incorrect.lower() in excluded_terms or re.match(r"[^a-zA-Z0-9]", incorrect):  
            grammar_issues.append({"incorrect": incorrect, "corrected": "N/A"})
            continue

        corrected = match.replacements[0] if match.replacements else "N/A"
        grammar_issues.append({"incorrect": incorrect, "corrected": corrected})

    return grammar_issues if grammar_issues else "No issues found"

# Resume length check
def check_length(text):
    words = text.split()
    return "Within recommended length" if 400 <= len(words) <= 800 else "Too short or too long"

# Long bullet points check
def check_bullet_length(text):
    bullets = re.findall(r"•\s*(.+)|-\s*(.+)|\\s(.+)", text)
    long_bullets = [b[0] for b in bullets if len(b[0].split()) > 20]
    return long_bullets if long_bullets else "All bullet points are fine"

# Contact info check
def check_contact_info(text):
    email = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone = re.search(r"(\+?\d{1,3})?[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}", text)
    linkedin = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9-]+", text)

    return {
        "email": bool(email),
        "phone": bool(phone),
        "linkedin": bool(linkedin),
        "missing": [field for field, present in {"email": email, "phone": phone, "linkedin": linkedin}.items() if not present]
    }

# Hard skills extraction
# def extract_hard_skills(text):
#     return [skill for skill in HARD_SKILLS if skill.lower() in text.lower()]

import pandas as pd
import os

import spacy
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
import re
import os
from collections import Counter

import spacy
import pandas as pd
import re
import os
from collections import Counter
import string

def extract_hard_skills(text, custom_skills_path=None):

    if not isinstance(text, str):
        return []
        
    # Load spaCy model for NLP processing
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # If model isn't available, download it
        import subprocess
        print("Downloading required NLP model...")
        subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
        nlp = spacy.load("en_core_web_sm")
    
    # Load skill dictionaries
    hard_skills_dict = _load_skills_dictionary(custom_skills_path)
    
    # Clean the input text
    text = _clean_text(text)
    
    # Process with NLP
    doc = nlp(text)
    
    # Extract potential skills using different methods
    extracted_skills = []
    
    # Method 1: Known hard skills direct matching
    for skill in hard_skills_dict:
        # Skip very short skills to avoid false positives
        if len(skill) < 3:
            continue
            
        # Use word boundary regex to find whole words/phrases
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            extracted_skills.append(skill)
    
    # Method 2: Extract noun phrases as potential skills
    for chunk in doc.noun_chunks:
        phrase = chunk.text.lower()
        # Check if the noun phrase is a hard skill or contains hard skills
        if _is_technical_phrase(phrase, hard_skills_dict):
            # Extract the main skill from the phrase
            clean_phrase = _clean_skill_phrase(phrase)
            if clean_phrase:
                extracted_skills.append(clean_phrase)
    
    # Method 3: Extract named entities that could be technologies
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PRODUCT"]:
            ent_text = ent.text.lower()
            if _is_technical_term(ent_text, hard_skills_dict):
                clean_ent = _clean_skill_phrase(ent_text)
                if clean_ent:
                    extracted_skills.append(clean_ent)
    
    # Method 4: Find known programming languages and technologies
    tech_terms = _extract_tech_terms(text)
    extracted_skills.extend(tech_terms)
    
    # Clean and normalize the extracted skills
    normalized_skills = _normalize_skills(extracted_skills)
    
    # Remove skills that are subsets of other skills
    final_skills = _remove_subset_skills(normalized_skills)
    
    # Sort alphabetically for consistency
    return sorted(final_skills)

def _load_skills_dictionary(custom_path=None):
    """
    Load a dictionary of hard skills. Attempts to load from custom path first,
    then falls back to default built-in dictionary.
    
    Returns:
        set: Set of hard skill terms
    """
    # Core technical skills dictionary
    HARD_SKILLS = {
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 
        'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'bash', 'shell', 'powershell',
        'perl', 'objective-c', 'sql', 'pl/sql', 't-sql',
        
        # Web Development
        'html', 'css', 'react', 'vue', 'angular', 'node.js', 'express', 'django',
        'flask', 'ruby on rails', 'spring boot', 'asp.net', 'laravel', 'jquery',
        'bootstrap', 'sass', 'less', 'webpack', 'graphql', 'rest api',
        
        # Mobile Development
        'android', 'ios', 'swift', 'react native', 'flutter', 'kotlin', 'xamarin',
        'ionic', 'cordova', 'objective-c',
        
        # Databases
        'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server', 'sqlite', 'redis',
        'cassandra', 'dynamodb', 'mariadb', 'neo4j', 'couchdb', 'firebase',
        
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins',
        'terraform', 'ansible', 'circleci', 'travis ci', 'github actions', 'gitlab ci',
        'cloudformation', 'nginx', 'apache', 'linux', 'unix',
        
        # Data Science & ML
        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy',
        'scipy', 'matplotlib', 'seaborn', 'tableau', 'power bi', 'machine learning',
        'deep learning', 'nlp', 'computer vision', 'data mining', 'hadoop', 'spark',
        
        # Version Control
        'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial',
        
        # Project Management
        'jira', 'confluence', 'agile', 'scrum', 'kanban', 'waterfall', 'prince2',
        'pmp', 'ms project', 'basecamp', 'asana', 'trello',
        
        # Design & Creative
        'photoshop', 'illustrator', 'indesign', 'sketch', 'figma', 'adobe xd',
        'after effects', 'premiere pro', 'blender', 'maya', '3ds max',
        
        # Office & Productivity
        'excel', 'word', 'powerpoint', 'access', 'vba', 'outlook', 'sharepoint',
        'sap', 'salesforce', 'zoho',
        
        # Other Technical Skills
        'rest api', 'soap', 'json', 'xml', 'yaml', 'regex', 'ui/ux', 'seo', 'devops',
        'ci/cd', 'blockchain', 'cryptography', 'networking', 'security', 'a/b testing'
    }
    
    # Try to load custom skills file if provided
    if custom_path and os.path.exists(custom_path):
        try:
            skills_df = pd.read_csv(custom_path)
            if 'skill' in skills_df.columns and 'type' in skills_df.columns:
                for _, row in skills_df.iterrows():
                    skill = str(row['skill']).lower().strip()
                    skill_type = str(row['type']).lower().strip()
                    if skill and skill_type == 'hard':
                        HARD_SKILLS.add(skill)
        except Exception as e:
            print(f"Error loading custom skills file: {e}")
    
    return HARD_SKILLS

def _clean_text(text):
    """Clean the input text for better processing"""
    # Replace problematic characters
    text = text.replace('●', ' ')
    text = text.replace(':-', ' ')
    text = text.replace('•', ' ')
    
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    return text.lower().strip()

def _is_technical_phrase(phrase, hard_skills_dict):
    """Check if a phrase is or contains a technical term"""
    phrase = phrase.lower()
    
    # Direct match
    if phrase in hard_skills_dict:
        return True
    
    # Check if any known hard skill is contained in the phrase
    for skill in hard_skills_dict:
        if len(skill) > 2 and re.search(r'\b' + re.escape(skill) + r'\b', phrase):
            return True
    
    # Check for technical indicators
    technical_indicators = ['software', 'programming', 'code', 'development', 'framework', 
                           'language', 'database', 'system', 'technology', 'api']
    
    for indicator in technical_indicators:
        if indicator in phrase:
            return True
            
    return False

def _is_technical_term(term, hard_skills_dict):
    """Determine if a term is a technical term based on dictionary and heuristics"""
    term = term.lower()
    
    # Direct match with hard skills dictionary
    if term in hard_skills_dict:
        return True
        
    # Check for version numbers (e.g., Python 3.8)
    if re.search(r'[a-zA-Z]+\s+\d+(\.\d+)*', term):
        return True
        
    return False

def _extract_tech_terms(text):
    """Extract known technical terms using regex patterns"""
    # Common programming languages, frameworks, and tools
    tech_pattern = r'\b(python|java|javascript|typescript|c\+\+|c#|ruby|php|swift|kotlin|go|' + \
                 r'react|angular|vue|node\.js|express|django|flask|spring boot|' + \
                 r'aws|azure|gcp|docker|kubernetes|jenkins|terraform|' + \
                 r'mysql|postgresql|mongodb|oracle|sql server|sqlite|redis|' + \
                 r'git|github|gitlab|bitbucket|' + \
                 r'tensorflow|pytorch|keras|pandas|numpy|scikit-learn)\b'
                 
    matches = set(re.findall(tech_pattern, text.lower()))
    return list(matches)

def _clean_skill_phrase(phrase):
    """Clean a skill phrase by removing fluff and keeping the core skill"""
    phrase = phrase.lower().strip()
    
    # Remove common fluff words that don't add skill value
    fluff_words = ['knowledge of', 'experience in', 'proficient in', 'skilled in', 
                  'expertise in', 'familiar with', 'working with', 'using', 
                  'ability to', 'skills in', 'including', 'such as', 'etc']
    
    for fluff in fluff_words:
        phrase = phrase.replace(fluff, ' ')
    
    # Remove punctuation except hyphens and periods (for terms like "node.js")
    for char in string.punctuation:
        if char not in ['-', '.']:
            phrase = phrase.replace(char, ' ')
    
    # Replace multiple spaces with single space
    phrase = re.sub(r'\s+', ' ', phrase).strip()
    
    # Don't return very short phrases (likely not meaningful skills)
    if len(phrase) < 3 or phrase in ['a', 'an', 'the', 'and', 'or']:
        return None
        
    return phrase

def _normalize_skills(skills_list):
    """Normalize skill names to avoid duplicates with different formatting"""
    normalized = []
    
    for skill in skills_list:
        if not skill:
            continue
            
        # Clean the skill name
        skill = skill.lower().strip()
        
        # Remove fluff words
        skill = re.sub(r'\b(knowledge|experience|proficient|skilled|expertise)\b', '', skill)
        
        # Replace separators with spaces
        skill = re.sub(r'[/,;:|]', ' ', skill)
        
        # Replace multiple spaces
        skill = re.sub(r'\s+', ' ', skill).strip()
        
        if skill and len(skill) > 2:
            normalized.append(skill)
    
    return list(set(normalized))

def _remove_subset_skills(skills_list):
    """Remove skills that are subsets of other skills in the list"""
    if not skills_list:
        return []
        
    # Sort by length (descending) to process longer phrases first
    sorted_skills = sorted(skills_list, key=len, reverse=True)
    final_skills = []
    
    for skill in sorted_skills:
        # Skip skills that are subsets of already accepted skills
        is_subset = False
        for accepted_skill in final_skills:
            # Check if this skill is contained within another skill
            if skill != accepted_skill and f" {skill} " in f" {accepted_skill} ":
                is_subset = True
                break
                
        # Avoid duplicating skills that are nearly identical
        similar_exists = False
        for accepted_skill in final_skills:
            # Check for similar skills (e.g., "react" vs "reactjs")
            if (skill in accepted_skill or accepted_skill in skill) and \
               abs(len(skill) - len(accepted_skill)) <= 3:
                similar_exists = True
                break
        
        if not is_subset and not similar_exists:
            final_skills.append(skill)
    
    return final_skills

# Soft skills extraction
def extract_soft_skills(text):
    return [skill for skill in SOFT_SKILLS if skill.lower() in text.lower()]

# Detecting generic resume layout (based on number of tables)
import pdfplumber
import re
from collections import Counter
import numpy as np
from PIL import Image
import io
import cv2
import pytesseract
from sklearn.cluster import KMeans

def detect_generic_layout(pdf_path):
    """
    Analyzes a resume PDF to determine if it uses a generic/outdated template
    or a modern ATS-friendly layout.
    
    This function examines multiple aspects of the resume:
    - Font variety and consistency
    - Layout structure and white space usage
    - Content-to-decoration ratio
    - Header and section formatting
    - Use of graphics and design elements
    - Modern vs outdated formatting patterns
    
    Args:
        pdf_path (str): Path to the resume PDF file
        
    Returns:
        dict: Analysis results including template classification and detailed metrics
    """
    results = {
        "classification": None,
        "confidence_score": 0,
        "metrics": {},
        "issues": [],
        "strengths": []
    }
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Skip analysis if PDF has too many pages (likely not a resume)
            if len(pdf.pages) > 10:
                results["classification"] = "Not a standard resume (too many pages)"
                results["issues"].append("Document exceeds typical resume length")
                return results
                
            # Extract text, fonts, and layout information
            all_text = ""
            fonts = []
            text_blocks = []
            graphics_count = 0
            headers = []
            line_spacing_values = []
            section_titles = []
            bullet_points = 0
            
            # Analyze each page
            for i, page in enumerate(pdf.pages):
                # Extract page text
                page_text = page.extract_text() or ""
                all_text += page_text
                
                # Analyze text characteristics
                if page.chars:
                    # Extract font information
                    page_fonts = [char.get('fontname', '') for char in page.chars]
                    fonts.extend(page_fonts)
                    
                    # Identify potential headers based on font size
                    font_sizes = [char.get('size', 0) for char in page.chars]
                    if font_sizes:
                        avg_font_size = sum(font_sizes) / len(font_sizes)
                        
                        # Group text by font size for header detection
                        text_by_size = {}
                        current_text = ""
                        current_size = None
                        
                        for j, char in enumerate(page.chars):
                            size = char.get('size', 0)
                            text = char.get('text', '')
                            
                            if current_size is None:
                                current_size = size
                                
                            if size == current_size:
                                current_text += text
                            else:
                                if current_text.strip():
                                    if current_size not in text_by_size:
                                        text_by_size[current_size] = []
                                    text_by_size[current_size].append(current_text.strip())
                                current_text = text
                                current_size = size
                                
                        # Add the last text segment
                        if current_text.strip() and current_size is not None:
                            if current_size not in text_by_size:
                                text_by_size[current_size] = []
                            text_by_size[current_size].append(current_text.strip())
                            
                        # Identify headers (larger font sizes)
                        largest_sizes = sorted(text_by_size.keys(), reverse=True)[:3]  # Top 3 largest font sizes
                        for size in largest_sizes:
                            if size > avg_font_size * 1.2:  # At least 20% larger than average
                                headers.extend(text_by_size[size])
                
                # Extract text blocks and their positions for layout analysis
                if page.extract_words():
                    words = page.extract_words()
                    
                    # Estimate line spacing
                    y_positions = [word['top'] for word in words]
                    y_positions.sort()
                    
                    if len(y_positions) > 1:
                        diffs = [y_positions[i+1] - y_positions[i] for i in range(len(y_positions) - 1)]
                        # Filter out large gaps (likely section breaks)
                        valid_diffs = [d for d in diffs if d < 20]
                        if valid_diffs:
                            line_spacing_values.extend(valid_diffs)
                    
                    # Analyze text block structure
                    for word in words:
                        text_blocks.append((word['text'], word['top'], word['bottom'], word['x0'], word['x1']))
                
                # Count bullet points
                bullet_patterns = [r'•', r'\\u2022', r'\\u25CF', r'\\u25CB', r'\\u25AA', r'\\u2043', 
                                  r'\\u2219', r'-\\s', r'\\*\\s', r'o\\s', r'\\d+\\.']
                for pattern in bullet_patterns:
                    bullet_points += len(re.findall(pattern, page_text))
                
                # Extract potential section titles
                section_patterns = [
                    r'^[A-Z][A-Z\s]+:',  # ALL CAPS with colon
                    r'^[A-Z][A-Z\s]+$',  # ALL CAPS line
                    r'^[A-Za-z\s]+:',    # Title with colon
                    r'EXPERIENCE|EDUCATION|SKILLS|PROJECTS|CERTIFICATIONS|SUMMARY|OBJECTIVE'
                ]
                
                for line in page_text.split('\n'):
                    for pattern in section_patterns:
                        if re.search(pattern, line.strip()):
                            section_titles.append(line.strip())
                            break
                
                # Count graphical elements (approximate)
                if page.images:
                    graphics_count += len(page.images)
                
                # Estimate decorative elements from curves, rects, and lines
                graphics_count += len(page.curves) + len(page.rects) + len(page.lines)
            
            # Calculate metrics
            metrics = {}
            
            # 1. Font consistency analysis
            font_counts = Counter(fonts)
            metrics["unique_fonts"] = len(font_counts)
            metrics["primary_font_ratio"] = max(font_counts.values()) / len(fonts) if fonts else 0
            
            # 2. Content density and organization
            metrics["text_length"] = len(all_text)
            metrics["bullet_points"] = bullet_points
            metrics["section_count"] = len(set(section_titles))
            
            # 3. Layout structure
            metrics["graphics_count"] = graphics_count
            metrics["content_to_decoration_ratio"] = len(all_text) / (graphics_count + 1)  # Avoid div by zero
            
            if line_spacing_values:
                metrics["avg_line_spacing"] = sum(line_spacing_values) / len(line_spacing_values)
                metrics["line_spacing_consistency"] = np.std(line_spacing_values) if len(line_spacing_values) > 1 else 0
            
            # 4. Calculate vertical whitespace distribution
            if text_blocks:
                # Sort blocks by vertical position
                sorted_blocks = sorted(text_blocks, key=lambda x: x[1])
                vertical_gaps = []
                
                for i in range(len(sorted_blocks) - 1):
                    current_block_bottom = sorted_blocks[i][2]
                    next_block_top = sorted_blocks[i+1][1]
                    gap = next_block_top - current_block_bottom
                    if gap > 0:
                        vertical_gaps.append(gap)
                
                if vertical_gaps:
                    metrics["whitespace_distribution"] = np.std(vertical_gaps) / np.mean(vertical_gaps) if np.mean(vertical_gaps) > 0 else 0
            
            # 5. Column detection (estimate through text block positions)
            if text_blocks:
                x_positions = [block[3] for block in text_blocks]  # x0 positions
                
                # Use k-means to detect potential columns (1-3 columns)
                max_columns = min(3, len(set(x_positions)))
                
                if max_columns > 1 and len(x_positions) > max_columns:
                    x_positions_array = np.array(x_positions).reshape(-1, 1)
                    
                    # Try different numbers of columns
                    best_score = -1
                    best_n_columns = 1
                    
                    for n_columns in range(1, max_columns + 1):
                        kmeans = KMeans(n_clusters=n_columns, random_state=0).fit(x_positions_array)
                        score = kmeans.inertia_
                        
                        # Lower score is better, but we penalize more columns
                        adjusted_score = score * (1 + (n_columns - 1) * 0.2)
                        
                        if best_score == -1 or adjusted_score < best_score:
                            best_score = adjusted_score
                            best_n_columns = n_columns
                    
                    metrics["estimated_columns"] = best_n_columns
                else:
                    metrics["estimated_columns"] = 1
            
            # Save all metrics
            results["metrics"] = metrics
            
            # Analyze and classify template
            # Calculate template modernity score (0-100)
            modernity_score = 0
            
            # Font consistency (modern resumes typically use 1-2 fonts consistently)
            if metrics["unique_fonts"] <= 2:
                modernity_score += 15
                results["strengths"].append("Appropriate font selection (1-2 fonts)")
            elif metrics["unique_fonts"] > 4:
                modernity_score -= 15
                results["issues"].append("Too many different fonts (more than 4)")
            
            # Primary font usage should be consistent
            if metrics["primary_font_ratio"] > 0.8:
                modernity_score += 10
                results["strengths"].append("Consistent primary font usage")
            elif metrics["primary_font_ratio"] < 0.6:
                modernity_score -= 10
                results["issues"].append("Inconsistent font usage")
            
            # Content to decoration ratio (modern resumes favor content over decoration)
            if metrics["content_to_decoration_ratio"] > 100:
                modernity_score += 15
                results["strengths"].append("Content-focused design (minimal decoration)")
            elif metrics["content_to_decoration_ratio"] < 50 and metrics["graphics_count"] > 5:
                modernity_score -= 15
                results["issues"].append("Too many decorative elements")
            
            # Column structure (1-2 columns is modern, 3+ columns is outdated)
            if metrics.get("estimated_columns", 1) <= 2:
                modernity_score += 10
                results["strengths"].append("Modern column structure")
            else:
                modernity_score -= 10
                results["issues"].append("Multi-column layout may cause ATS issues")
            
            # Section organization
            if metrics["section_count"] >= 4:
                modernity_score += 10
                results["strengths"].append("Well-organized with clear sections")
            elif metrics["section_count"] < 3:
                modernity_score -= 5
                results["issues"].append("Limited section organization")
            
            # Bullet point usage (modern resumes use bullet points effectively)
            if 5 <= metrics["bullet_points"] <= 30:
                modernity_score += 10
                results["strengths"].append("Effective use of bullet points")
            elif metrics["bullet_points"] > 40:
                modernity_score -= 5
                results["issues"].append("Excessive bullet points")
            elif metrics["bullet_points"] < 3:
                results["issues"].append("Limited use of bullet points")
            
            # Whitespace distribution (modern resumes use whitespace strategically)
            if "whitespace_distribution" in metrics:
                if metrics["whitespace_distribution"] < 0.5:
                    modernity_score += 10
                    results["strengths"].append("Strategic use of whitespace")
                elif metrics["whitespace_distribution"] > 1.0:
                    modernity_score -= 10
                    results["issues"].append("Inconsistent spacing")
            
            # Line spacing consistency
            if "line_spacing_consistency" in metrics:
                if metrics["line_spacing_consistency"] < 2:
                    modernity_score += 10
                    results["strengths"].append("Consistent line spacing")
                elif metrics["line_spacing_consistency"] > 5:
                    modernity_score -= 10
                    results["issues"].append("Inconsistent line spacing")
            
            # Determine classification based on score
            results["modernity_score"] = modernity_score
            
            if modernity_score >= 60:
                results["classification"] = "Modern (ATS-friendly layout)"
                results["confidence_score"] = (modernity_score - 60) / 40  # 0.0 to 1.0
            elif modernity_score >= 40:
                results["classification"] = "Average (Some modern elements)"
                results["confidence_score"] = (modernity_score - 40) / 20  # 0.0 to 1.0
            else:
                results["classification"] = "Generic/Outdated (Not ATS-optimized)"
                results["confidence_score"] = (40 - modernity_score) / 40  # 0.0 to 1.0
            
    except Exception as e:
        results["classification"] = "Error analyzing template"
        results["issues"].append(f"Analysis error: {str(e)}")
    
    return results

# Main execution

file_path = os.path.expanduser(r"C:\Users\Pradeep babu\Downloads\PHANI DATTA KANDUKURI.pdf")  # Adjust path for Windows if needed

text = extract_text(file_path)

# Use the new corrected parse rate function
parse_rate = calculate_parse_rate(file_path, text)
impact = quantify_impact(text)
repeated_words = check_repetition(text)
grammar_issues = check_grammar(text)
format_check = check_file_format(file_path)
resume_length = check_length(text)
long_bullets = check_bullet_length(text)
contact_info = check_contact_info(text)
# sections = check_sections(text)
hard_skills = extract_hard_skills(text)
soft_skills = extract_soft_skills(text)
design_type = detect_generic_layout(file_path)

ats_score = (
    min(parse_rate, 100) * 0.35 +  # Normalize parse rate within 100
    (5 if impact else 0) +  
    (-5 if repeated_words != "No excessive repetition" else 5) +  
    max(0, 35 - len(grammar_issues) * 2) * 0.35 +  
    (5 if format_check == "Valid (PDF format)" else 0) +  
    (5 if resume_length == "Within recommended length" else 0) +  
    (5 if long_bullets == "All bullet points are fine" else 0) +  
    (10 if len(contact_info["missing"]) == 0 else 5) +  
    min(len(hard_skills) * 2, 10) +  
    min(len(soft_skills) * 2, 10) +  
    (10 if design_type == "Unique (ATS-friendly layout)" else 0)
)

print(text)

print({
    "ats_score": round(ats_score, 2),
    "parse_rate": parse_rate,
    "impact": impact,
    "repeated_words": repeated_words,
    "grammar_issues": grammar_issues,
    "format_check": format_check,
    "resume_length": resume_length,
    "long_bullets": long_bullets,
    "contact_info": contact_info,
    # "sections": sections,
    "hard_skills": hard_skills,
    "soft_skills": soft_skills,
    "design_type": design_type
})


from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import os
import shutil
import PyPDF2
import docx
import re

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins (change this for security)
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)

@app.get("/")
def read_root():
    return {"message": "Resume Analysis API is running"}

def extract_text(file_path):
    """
    Extract text from PDF or DOCX files
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    elif file_ext in ['.docx', '.doc']:
        try:
            if file_ext == '.docx':
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            else:
                # For .doc files, you might need additional libraries
                # This is a placeholder - you may need to implement conversion
                return "DOC format not fully supported yet"
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            return ""
    
    else:
        return "Unsupported file format"

def calculate_parse_rate(file_path, text):
    """
    Calculate how well the resume was parsed
    Here's a simple implementation that evaluates based on content identification
    """
    if not text or len(text) < 10:
        return 0.0
    
    # Check for common resume sections
    sections = {
        'contact': re.search(r'(email|phone|address|linkedin)', text.lower()) is not None,
        'education': re.search(r'(education|university|college|degree|bachelor|master|phd)', text.lower()) is not None,
        'experience': re.search(r'(experience|work|employment|job|position|company)', text.lower()) is not None,
        'skills': re.search(r'(skills|abilities|proficiencies|technologies|languages)', text.lower()) is not None
    }
    
    # Calculate parse rate based on sections identified
    identified_sections = sum(sections.values())
    total_sections = len(sections)
    
    # Basic parse rate calculation
    parse_rate = (identified_sections / total_sections) * 100
    
    # Adjust parse rate based on text length (longer text might mean better extraction)
    text_quality_bonus = min(len(text) / 1000, 20)  # Max 20% bonus for text length
    
    final_rate = min(parse_rate + text_quality_bonus, 100)
    return round(final_rate, 2)

from fastapi import FastAPI, UploadFile, File, Form
import os
import json
from pydantic import BaseModel

app = FastAPI()

class JobSearchRequest(BaseModel):
    location_pref: str
    work_mode: str
    limit: int = 10

@app.post("/analyze_resume")
async def analyze_resume(
    file: UploadFile = File(...), 
    location_pref: str = Form("India"), 
    work_mode: str = Form("all"), 
    limit: int = Form(10)
):
    file_path = f"temp_{file.filename}"
    with open(file_path, "wb") as f:
        f.write(await file.read())
    
    # Extract text from resume
    text = extract_text(file_path)

    # Perform ATS scoring
    parse_rate = calculate_parse_rate(file_path, text)
    impact = quantify_impact(text)
    repeated_words = check_repetition(text)
    grammar_issues = check_grammar(text)
    format_check = check_file_format(file_path)
    resume_length = check_length(text)
    long_bullets = check_bullet_length(text)
    contact_info = check_contact_info(text)
    hard_skills = extract_hard_skills(text)
    soft_skills = extract_soft_skills(text)
    design_type = detect_generic_layout(file_path)
    all_skills = list(set(hard_skills + soft_skills))

    ats_score = (
        min(parse_rate, 100) * 0.35 +
        (5 if impact else 0) +
        (-5 if repeated_words != "No excessive repetition" else 5) +
        max(0, 35 - len(grammar_issues) * 2) * 0.35 +
        (5 if format_check == "Valid (PDF format)" else 0) +
        (5 if resume_length == "Within recommended length" else 0) +
        (5 if long_bullets == "All bullet points are fine" else 0) +
        (10 if len(contact_info["missing"]) == 0 else 5) +
        min(len(hard_skills) * 2, 10) +
        min(len(soft_skills) * 2, 10) +
        (10 if design_type == "Unique (ATS-friendly layout)" else 0)
    )

    # Determine job location preference
  
    
    # Fetch job listings

    return {
        "ats_score": round(ats_score, 2),
        "parse_rate": parse_rate,
        "impact": impact,
        "repeated_words": repeated_words,
        "grammar_issues": grammar_issues,
        "format_check": format_check,
        "resume_length": resume_length,
        "long_bullets": long_bullets,
        "contact_info": contact_info,
        "hard_skills": hard_skills,
        "soft_skills": soft_skills,
        "design_type": design_type,
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5009)